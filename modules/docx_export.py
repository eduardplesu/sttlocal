from docx import Document

def ticks_to_time(ticks):
    """
    Converts ticks (100-nanosecond intervals) to a formatted time string HH:MM:SS.mmm.
    """
    seconds = ticks / 10_000_000  # convert ticks to seconds
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = seconds % 60
    return f"{hours:02d}:{minutes:02d}:{secs:06.3f}"

def export_transcription_to_docx(transcription_results, analysis_text=None, 
                                 translated_transcription=None, cleaned_transcription=None, 
                                 output_filename="transcription.docx"):
    """
    Exports transcription results to a DOCX file.

    The DOCX file will include:
      - Transcription segments with speaker information and time details.
      - An optional Cleaned Transcription section.
      - An optional Translated Transcription section.
      - An optional Analysis section.

    Parameters:
      - transcription_results (list): List of dictionaries for each transcription segment.
      - analysis_text (str): Optional analysis text to include.
      - translated_transcription (str): Optional translated transcription text.
      - cleaned_transcription (str): Optional cleaned transcription text.
      - output_filename (str): Output file name for the DOCX document.
    
    Returns:
      - str: The output file name of the generated DOCX file.
    """
    document = Document()
    document.add_heading("Transcription", level=0)
    
    # Add transcription segments with speaker and timing details.
    for result in transcription_results:
        speaker = result.get("speaker_name", result.get("speaker_id", "Unknown"))
        text = result.get("text", "")
        offset = result.get("offset", 0)
        duration = result.get("duration", 0)
        start_time = ticks_to_time(offset)
        duration_time = ticks_to_time(duration)
        document.add_paragraph(f"Speaker {speaker}: {text}")
        document.add_paragraph(f"(Start Time: {start_time}, Duration: {duration_time})", style="Intense Quote")
    
    # Add Cleaned Transcription section if provided.
    if cleaned_transcription:
        document.add_page_break()
        document.add_heading("Cleaned Transcription", level=0)
        document.add_paragraph(cleaned_transcription)
    
    # Add Translated Transcription section if provided.
    if translated_transcription:
        document.add_page_break()
        document.add_heading("Translated Transcription", level=0)
        document.add_paragraph(translated_transcription)
    
    # Add Analysis section if provided.
    if analysis_text:
        document.add_page_break()
        document.add_heading("Analysis", level=0)
        document.add_paragraph(analysis_text)
    
    document.save(output_filename)
    return output_filename
